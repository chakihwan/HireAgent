"""파일 로더 — PDF / DOCX / MD / TXT 텍스트 추출."""
import io

import docx
import pypdf


class FileParseError(Exception):
    """파일 파싱 실패 시 사용자 표시용 메시지."""


_MAX_SIZE_BYTES = 20 * 1024 * 1024  # 20MB


def parse_file(filename: str, data: bytes) -> str:
    """업로드된 파일 바이트 → 텍스트 추출.

    지원: pdf, docx, md, markdown, txt
    실패 시 FileParseError.
    """
    if len(data) > _MAX_SIZE_BYTES:
        raise FileParseError(f"파일이 너무 큽니다 ({len(data) // (1024 * 1024)}MB > 20MB 제한).")

    name = filename.lower()
    try:
        if name.endswith(".pdf"):
            return _parse_pdf(data)
        if name.endswith(".docx"):
            return _parse_docx(data)
        if name.endswith((".md", ".markdown", ".txt")):
            return _parse_text(data)
    except FileParseError:
        raise
    except Exception as e:
        raise FileParseError(f"파일 파싱 실패: {e}")

    raise FileParseError(
        f"지원하지 않는 파일 형식입니다: {filename}. PDF/DOCX/MD/TXT만 지원합니다."
    )


def _parse_pdf(data: bytes) -> str:
    reader = pypdf.PdfReader(io.BytesIO(data))
    if reader.is_encrypted:
        raise FileParseError("암호화된 PDF는 지원하지 않습니다.")
    pages = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        if text.strip():
            pages.append(text.strip())
    if not pages:
        raise FileParseError(
            "PDF에서 텍스트를 추출하지 못했습니다. 이미지 기반 PDF일 수 있습니다."
        )
    return "\n\n".join(pages)


def _parse_docx(data: bytes) -> str:
    doc = docx.Document(io.BytesIO(data))
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    # 표 안의 텍스트도 추출
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)
    if not parts:
        raise FileParseError("DOCX에서 텍스트를 추출하지 못했습니다.")
    return "\n".join(parts)


def _parse_text(data: bytes) -> str:
    # UTF-8 우선, 실패 시 CP949 (한국어 인코딩 대응)
    for encoding in ("utf-8", "utf-8-sig", "cp949", "euc-kr"):
        try:
            return data.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    raise FileParseError("텍스트 인코딩을 인식할 수 없습니다 (UTF-8/CP949).")
